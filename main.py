# -*- coding: utf-8 -*-
"""
IAttom – WhatsApp Bot (Render/Flask) com:
- Conversa natural (LLM fallback) + persona empática
- Comandos: ajuda, foco, diário, imagem, pdf, docx, wiki, pesquisar, resumo, reset
- Memória leve (nome, diário)
- Geração de arquivos (PDF/DOCX) e entrega via link/arquivo
- Servidor de arquivos /files/<nome>
"""

import os, re, time, json
from datetime import datetime
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from flask import Flask, request, send_from_directory

# ===================== 0) CONFIG & VARS =====================

ACCESS_TOKEN    = os.getenv("ACCESS_TOKEN",    "SEU_TOKEN_AQUI")
PHONE_NUMBER_ID = os.getenv("PHONE_NUMBER_ID", "730613666807699")
VERIFY_TOKEN    = os.getenv("VERIFY_TOKEN",    "IATTOM2025")
BUSINESS_ID     = os.getenv("BUSINESS_ACCOUNT_ID", "1090223772743358")

OPENAI_API_KEY  = os.getenv("OPENAI_API_KEY", "")      # para cérebro IA e imagens
APP_BASE_URL    = os.getenv("APP_BASE_URL", "")        # ex.: https://iattom-bot.onrender.com
SERPAPI_KEY     = os.getenv("SERPAPI_KEY", "")         # opcional: pesquisa web melhor
EMO_CHECKIN_HOURS = int(os.getenv("EMO_CHECKIN_HOURS", "6"))

GRAPH_URL = f"https://graph.facebook.com/v23.0/{PHONE_NUMBER_ID}/messages"
HEADERS   = {"Authorization": f"Bearer {ACCESS_TOKEN}", "Content-Type": "application/json"}

# Arquivos gerados
FILES_DIR = Path("/tmp/files")
FILES_DIR.mkdir(parents=True, exist_ok=True)

# Memória leve
PROFILE = {}         # wa_id -> dict(name,last_hi,last_emo,diary[])
FIRST_HEART = "💜"   # só em aberturas / apoio

# ===================== 1) HELPERS WHATSAPP =====================

def send_text(wa_id: str, body: str):
    payload = {"messaging_product":"whatsapp","to":wa_id,"type":"text","text":{"body":body}}
    r = requests.post(GRAPH_URL, headers=HEADERS, json=payload, timeout=30)
    print("➡️ send_text:", r.status_code, r.text[:500])
    return r.ok

def send_image_link(wa_id: str, image_url: str, caption: str = ""):
    payload = {"messaging_product":"whatsapp","to":wa_id,"type":"image",
               "image":{"link":image_url,"caption":caption[:1024]}}
    r = requests.post(GRAPH_URL, headers=HEADERS, json=payload, timeout=30)
    print("➡️ send_image_link:", r.status_code, r.text[:500])
    return r.ok

def send_document_link(wa_id: str, url: str, filename: str):
    payload = {
        "messaging_product": "whatsapp",
        "to": wa_id,
        "type": "document",
        "document": {"link": url, "filename": filename}
    }
    r = requests.post(GRAPH_URL, headers=HEADERS, json=payload, timeout=30)
    print("➡️ send_document_link:", r.status_code, r.text[:500])
    return r.ok

# ===================== 2) GERAÇÃO DE ARQUIVOS =====================

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

def safe_filename(name: str) -> str:
    name = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_")
    return name or "arquivo"

def create_docx(title: str, content: str) -> Path:
    doc = Document()
    doc.add_heading(title, 0)
    for para in content.split("\n"):
        doc.add_paragraph(para)
    fname = safe_filename(f"{title}.docx")
    fpath = FILES_DIR / fname
    doc.save(fpath)
    return fpath

def create_pdf(title: str, content: str) -> Path:
    fname = safe_filename(f"{title}.pdf")
    fpath = FILES_DIR / fname
    c = canvas.Canvas(str(fpath), pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, title[:90]); y -= 30
    c.setFont("Helvetica", 12)
    for line in content.split("\n"):
        if y < 50:
            c.showPage()
            y = height - 50
            c.setFont("Helvetica", 12)
        c.drawString(50, y, line[:100]); y -= 18
    c.save()
    return fpath

# ===================== 3) IA (LLM) & IMAGENS =====================

def llm_reply(system_prompt: str, user_prompt: str, max_tokens: int = 400) -> str:
    """
    Usa OpenAI Chat completions como 'cérebro'. Retorna texto.
    """
    if not OPENAI_API_KEY:
        return ""
    try:
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type":"application/json"},
            json={
                "model": "gpt-4o-mini",  # pode trocar por gpt-4o, gpt-4.1 etc.
                "messages": [
                    {"role":"system", "content": system_prompt},
                    {"role":"user",   "content": user_prompt}
                ],
                "temperature": 0.7,
                "max_tokens": max_tokens
            },
            timeout=60
        )
        if resp.status_code == 200:
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()
        print("⚠️ LLM error:", resp.status_code, resp.text[:500])
    except Exception as e:
        print("⚠️ LLM exception:", e)
    return ""

def generate_image_url(prompt: str) -> str:
    """
    Gera imagem (OpenAI Images). Retorna URL.
    """
    if not OPENAI_API_KEY:
        return ""
    try:
        r = requests.post(
            "https://api.openai.com/v1/images/generations",
            headers={"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type":"application/json"},
            json={"model":"gpt-image-1","prompt":prompt,"size":"1024x1024","n":1},
            timeout=60
        )
        if r.status_code == 200:
            data = r.json()
            return data["data"][0].get("url","") or ""
        print("⚠️ OpenAI image error:", r.status_code, r.text[:500])
    except Exception as e:
        print("⚠️ OpenAI image exception:", e)
    return ""

# ===================== 4) WIKI / PESQUISA / RESUMO =====================

def wiki_search(term: str) -> str:
    try:
        url = "https://pt.wikipedia.org/api/rest_v1/page/summary/" + requests.utils.quote(term.strip())
        r = requests.get(url, timeout=20)
        if r.status_code == 200:
            data = r.json()
            title = data.get("title","")
            extract = data.get("extract","")
            if extract:
                return f"{title}: {extract}"
    except Exception as e:
        print("⚠️ wiki exception:", e)
    return "Não encontrei um resumo claro na Wikipédia."

def ddg_search(query: str, max_results: int = 3):
    """
    Pesquisa rápida usando DuckDuckGo HTML (sem chave).
    Se tiver SERPAPI_KEY, você pode trocar por SerpAPI para resultados melhores.
    """
    try:
        if SERPAPI_KEY:
            url = "https://serpapi.com/search.json"
            r = requests.get(url, params={"engine":"google","q":query,"hl":"pt-br","num":5,"api_key":SERPAPI_KEY}, timeout=20)
            out=[]
            if r.status_code==200:
                js=r.json()
                for item in (js.get("organic_results") or [])[:max_results]:
                    out.append({"title": item.get("title",""), "link": item.get("link","")})
            return out if out else []
        # fallback ddg
        r = requests.get("https://duckduckgo.com/html/", params={"q":query}, timeout=20, headers={"User-Agent":"Mozilla/5.0"})
        soup = BeautifulSoup(r.text, "html.parser")
        results=[]
        for a in soup.select(".result__a")[:max_results]:
            title = a.get_text(" ", strip=True)
            link = a.get("href","")
            results.append({"title": title, "link": link})
        return results
    except Exception as e:
        print("⚠️ search exception:", e)
        return []

def summarize_url(url: str) -> str:
    """
    Baixa HTML e pede ao LLM um resumo em tópicos.
    """
    try:
        html = requests.get(url, timeout=20, headers={"User-Agent":"Mozilla/5.0"}).text
    except Exception as e:
        print("⚠️ fetch url exception:", e)
        return "Não consegui acessar o link para resumir."
    if not OPENAI_API_KEY:
        return "Consigo resumir o link quando você definir a OPENAI_API_KEY."
    text = BeautifulSoup(html, "html.parser").get_text(separator="\n")
    text = re.sub(r"\n{2,}", "\n", text).strip()
    prompt = f"Resuma de forma clara e objetiva, em 5 a 8 tópicos, o conteúdo a seguir.\n\nTexto:\n{text[:8000]}"
    sys = "Você é um assistente que cria resumos fiéis, concisos e úteis em português."
    return llm_reply(sys, prompt, max_tokens=350) or "Não consegui gerar um resumo agora."

# ===================== 5) MEMÓRIA & PERSONALIDADE =====================

def _now(): return time.time()

def mem_get(wa_id: str) -> dict:
    return PROFILE.get(wa_id, {})

def mem_set(wa_id: str, data: dict):
    PROFILE[wa_id] = data

def mem_reset(wa_id: str):
    PROFILE.pop(wa_id, None)

HELP_TEXT = (
    "Claro! Eu posso te ajudar com alguns comandos principais:\n\n"
    "✨ *Qual o seu nome?* — Eu digo quem eu sou.\n\n"
    "✨ *Ajuda* ou *Menu* — Eu mostro esta lista de opções.\n\n"
    "✨ *Foco* — Técnica Pomodoro: 25 minutos de concentração + 5 minutos de pausa.\n\n"
    "✨ *Imagem: descrição* — Eu gero uma imagem de acordo com o que você pedir (requer chave configurada).\n\n"
    "✨ *Diário: texto* — Eu guardo a sua anotação no seu diário pessoal.\n\n"
    "✨ *PDF: Título | conteúdo* — Eu crio um arquivo PDF e envio o link/arquivo.\n\n"
    "✨ *DOCX: Título | conteúdo* — Eu crio um arquivo do Word e envio o link/arquivo.\n\n"
    "✨ *Wiki: termo* — Eu busco um resumo da Wikipédia.\n\n"
    "✨ *Pesquisar: assunto* — Eu procuro na web e trago 2–3 fontes.\n\n"
    "✨ *Resumo: URL* — Eu leio o link e resumo em tópicos.\n\n"
    "✨ *Reset* — Eu limpo suas preferências e memórias para começar do zero.\n\n"
    "Pode escolher qualquer um desses que eu ajudo você. 😉"
)

NAME_PATTERNS = [
    r"(?:meu nome é|pode me chamar de|sou o|sou a)\s+([A-Za-zÀ-ÿ'´`^~\- ]{2,30})",
]

TRIGGERS_EMO  = ["triste","desanim","cansad","ansios","depress","sem vontade","sobrecarreg","estress"]
TRIGGERS_PROD = ["organizar","produtiv","prioridade","planejar","agenda","projeto","procrast"]
TRIGGERS_STUD = ["estudar","prova","concurso","enem","vestibular","matéria","resumo","memoriz"]

SIGNOFF = "E lembre-se: você é capaz de muito mais do que imagina."

def normalize_display_name(raw: str) -> str:
    if not raw: return ""
    raw = re.sub(r"^\s*(hoje|today)[^A-Za-zÀ-ÿ]*", "", raw, flags=re.I)
    raw = re.sub(r"[,\s]{2,}", " ", raw).strip()
    first = raw.split()[0] if raw else ""
    return first.title()

def extract_name_from_text(original: str) -> str|None:
    low = original.lower()
    for pat in NAME_PATTERNS:
        m = re.search(pat, low)
        if m:
            try:
                s, e = m.start(1), m.end(1)
                return original[s:e].strip().title()
            except:
                return m.group(1).strip().title()
    return None

def persona_reply(name: str|None, incoming: str) -> str:
    # Gatihos empatia/produtividade/estudos
    low = incoming.lower()
    if any(t in low for t in TRIGGERS_EMO):
        who = f"{name}, " if name else ""
        return (f"{who}eu tô com você. Respira comigo. Se quiser, me conta o que está pesando — "
                "a gente quebra em passos pequenos. " + FIRST_HEART)
    if any(t in low for t in TRIGGERS_PROD):
        who = f"{name}, " if name else ""
        return (f"{who}vamos por partes: 1) Liste 3 coisas que importam hoje. 2) Comece pela menor. "
                "3) Me avise quando concluir a primeira — sigo com você.")
    if any(t in low for t in TRIGGERS_STUD):
        who = f"{name}, " if name else ""
        return (f"{who}estudo rende mais com blocos curtos: 25min foco + 5min pausa (Pomodoro). "
                "Quer que eu te lembre?")

    who = f"{name}, " if name else ""
    return (f"{who}entendi. Me dá um pouquinho mais de contexto pra eu te ajudar melhor? "
            "Se for algo rápido, posso te sugerir o primeiro passo. " + FIRST_HEART)

def brain_fallback(name: str|None, incoming: str) -> str:
    """
    Cérebro IA: se não for comando, tenta responder bem.
    - Se ver URL única -> resumo do link
    - Se parecer termo enciclopédico -> wiki
    - Caso geral -> LLM com persona IAttom
    """
    # URL única?
    urls = re.findall(r"https?://\S+", incoming)
    if len(urls) == 1:
        return summarize_url(urls[0])

    # Curto e enciclopédico? tenta wiki
    if len(incoming.split()) <= 6 and not incoming.endswith("?"):
        w = wiki_search(incoming)
        if w and "Não encontrei" not in w:
            return w

    # LLM persona
    if not OPENAI_API_KEY:
        # Sem LLM, responde no tom da persona
        return persona_reply(name, incoming)

    system = (
        "Você é o IAttom, um assistente de IA empático, claro e útil. "
        "Responda em português, com tom humano, objetivo e acolhedor. "
        "Quando apropriado, proponha um primeiro passo concreto. "
        "Evite respostas longas demais; priorize clareza. "
        "Não invente fatos; se não tiver certeza, explique como podemos buscar juntos."
    )
    prefix = (f"Usuário: {name}\n" if name else "")
    user = prefix + incoming
    out = llm_reply(system, user, max_tokens=450)
    return out or persona_reply(name, incoming)

# ===================== 6) COMANDOS =====================

def handle_commands(wa_id: str, text: str) -> bool:
    low = text.strip().lower()

    # ajuda/menu
    if low in ("ajuda", "menu", "help"):
        send_text(wa_id, HELP_TEXT); return True

    # foco
    if low in ("foco", "pomodoro"):
        send_text(wa_id,
            "Vamos no simples: 25 minutos de foco + 5 minutos de pausa (Pomodoro).\n"
            "1) Escolha 1 tarefa pequena.\n"
            "2) Me avise quando concluir a primeira — sigo com você.")
        return True

    # identidade
    if low in ("qual o seu nome", "qual o seu nome?", "quem é você", "quem é voce", "quem e voce"):
        send_text(wa_id, "Eu sou o IAttom, o seu Assistente de Inteligência Artificial. 😊")
        return True

    # imagem: aceita “imagem:”, “img:”, “image:”
    if low.startswith(("imagem:", "img:", "image:")):
        prompt = text.split(":",1)[1].strip() if ":" in text else ""
        if not prompt:
            send_text(wa_id, "Diga o que você quer que eu gere: ex. *Imagem: um pôr do sol em aquarela*.")
            return True
        url = generate_image_url(prompt)
        if url: send_image_link(wa_id, url, caption=f"IAttom – imagem: {prompt}")
        else:   send_text(wa_id, "Consigo gerar imagens quando você definir sua OPENAI_API_KEY no Render.")
        return True

    # diário
    if low.startswith(("diário:", "diario:")):
        note = text.split(":",1)[1].strip() if ":" in text else ""
        p = mem_get(wa_id)
        diary = p.get("diary", [])
        if note:
            diary.append({"at": datetime.utcnow().isoformat()+"Z", "text": note})
            p["diary"] = diary; mem_set(wa_id, p)
            send_text(wa_id, "Anotado no seu diário. 📘")
        else:
            if diary:
                last = diary[-1]["text"]; send_text(wa_id, f"Sua última anotação foi: “{last}”.")
            else:
                send_text(wa_id, "Seu diário está vazio. Escreva: *Diário: ...*")
        return True

    # pdf / docx
    if low.startswith("pdf:"):
        raw = text[4:].strip()
        if "|" not in raw:
            send_text(wa_id, "Use assim: *PDF: Título | conteúdo*")
            return True
        title, body = [p.strip() for p in raw.split("|", 1)]
        path = create_pdf(title or "Arquivo", body or "")
        if not APP_BASE_URL:
            send_text(wa_id, "Defina APP_BASE_URL no Render (ex.: https://seuapp.onrender.com).")
            return True
        url = f"{APP_BASE_URL}/files/{path.name}"
        send_document_link(wa_id, url, path.name)
        send_text(wa_id, f"PDF gerado com sucesso:\n{url}")
        return True

    if low.startswith("docx:"):
        raw = text[5:].strip()
        if "|" not in raw:
            send_text(wa_id, "Use assim: *DOCX: Título | conteúdo*")
            return True
        title, body = [p.strip() for p in raw.split("|", 1)]
        path = create_docx(title or "Documento", body or "")
        if not APP_BASE_URL:
            send_text(wa_id, "Defina APP_BASE_URL no Render (ex.: https://seuapp.onrender.com).")
            return True
        url = f"{APP_BASE_URL}/files/{path.name}"
        send_document_link(wa_id, url, path.name)
        send_text(wa_id, f"DOCX criado com sucesso:\n{url}")
        return True

    # wiki / pesquisar / resumo
    if low.startswith("wiki:"):
        term = text.split(":",1)[1].strip() if ":" in text else ""
        send_text(wa_id, wiki_search(term or "")); return True

    if low.startswith("pesquisar:"):
        q = text.split(":",1)[1].strip() if ":" in text else ""
        results = ddg_search(q or "")
        if not results:
            send_text(wa_id, "Não achei resultados agora. Tente ser mais específico.")
            return True
        msg = "Aqui está o que encontrei:\n"
        for i, it in enumerate(results, 1):
            msg += f"{i}) {it.get('title','(sem título)')}\n{it.get('link','')}\n\n"
        send_text(wa_id, msg.strip()); return True

    if low.startswith("resumo:"):
        url = text.split(":",1)[1].strip() if ":" in text else ""
        send_text(wa_id, summarize_url(url or "")); return True

    # reset
    if low == "reset":
        mem_reset(wa_id); send_text(wa_id, "Pronto! Zerei suas preferências e memórias locais.")
        return True

    return False

# ===================== 7) FLASK ROUTES =====================

app = Flask(__name__)

@app.route("/", methods=["GET"])
def root_ok():
    return "IAttom online ✅", 200

@app.route("/files/<path:filename>", methods=["GET"])
def serve_files(filename):
    return send_from_directory(str(FILES_DIR), filename, as_attachment=False)

# Verificação do webhook
@app.route("/webhook", methods=["GET"])
def verify():
    mode      = request.args.get("hub.mode")
    token     = request.args.get("hub.verify_token")
    challenge = request.args.get("hub.challenge")
    if mode == "subscribe" and token == VERIFY_TOKEN:
        return challenge, 200
    return "Forbidden", 403

# Recebimento de mensagens
@app.route("/webhook", methods=["POST"])
def receive():
    data = request.get_json(silent=True, force=True)
    print("📥 evento:", json.dumps(data, ensure_ascii=False)[:1500])

    try:
        entry   = data.get("entry", [])[0]
        change  = entry.get("changes", [])[0]
        value   = change.get("value", {})
        msgs    = value.get("messages")
        if not msgs: return "OK", 200

        msg   = msgs[0]
        wa_id = msg["from"]

        # nome do contato (limpando “Hoje,, …” etc.)
        contacts = value.get("contacts", [])
        name = None
        if contacts and "profile" in contacts[0]:
            raw_name = contacts[0]["profile"].get("name", "")
            cleaned  = normalize_display_name(raw_name)
            if cleaned: 
                p = mem_get(wa_id); 
                if not p.get("name"): 
                    p["name"] = cleaned; mem_set(wa_id, p)
                name = p.get("name")
        else:
            p = mem_get(wa_id); name = p.get("name")

        if msg.get("type") == "text":
            body = msg["text"]["body"].strip()

            # aprender nome por texto (“meu nome é …”)
            new_name = extract_name_from_text(body)
            if new_name:
                p = mem_get(wa_id); p["name"] = new_name; mem_set(wa_id, p)
                send_text(wa_id, f"Prazer te conhecer, {new_name}! Pode contar comigo no que precisar."); 
                return "OK", 200

            # primeira saudação
            p = mem_get(wa_id)
            if "last_hi" not in p:
                p["last_hi"] = _now(); mem_set(wa_id, p)
                if not name:
                    send_text(wa_id, "Oi! Eu sou o IAttom — o seu Assistente de Inteligência Artificial. Como posso te chamar? " + FIRST_HEART)
                else:
                    send_text(wa_id, f"Oi, {name}! Que bom te ver por aqui. Em que posso te ajudar hoje? {FIRST_HEART}")
                # não retorna: deixa cair nos comandos ou cérebro

            # COMANDOS
            if handle_commands(wa_id, body):
                return "OK", 200

            # CHECK-IN emocional periódico
            last_emo = p.get("last_emo", 0.0)
            if (_now() - last_emo) > EMO_CHECKIN_HOURS * 3600:
                p["last_emo"] = _now(); mem_set(wa_id, p)
                send_text(wa_id, (f"{name+', ' if name else ''}como você tem se sentido hoje? Quero ter certeza de que você está bem. " + FIRST_HEART))

            # CÉREBRO IA (fallback)
            reply = brain_fallback(name, body)
            if reply.strip():
                # assinatura leve em respostas gerais
                if SIGNOFF not in reply:
                    reply = reply.strip() + "\n\n" + SIGNOFF
                send_text(wa_id, reply)

    except Exception as e:
        print("⚠️ erro no processamento:", e)

    return "OK", 200

# ===================== 8) RUN =====================

if __name__ == "__main__":
    port = int(os.getenv("PORT", "8080"))
    app.run(host="0.0.0.0", port=port)


